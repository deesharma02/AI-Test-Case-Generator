import streamlit as st
from dotenv import load_dotenv
from docx import Document
import io

# LangChain Imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_community.vectorstores import FAISS
from langchain_core.runnables import RunnableSequence
from langchain_core.documents import Document as LCDocument
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_classic.retrievers import MultiQueryRetriever
import uuid


# Load env
load_dotenv()

# -------------------------------
# Query Generation
# -------------------------------
@st.cache_data(show_spinner=False)
def generate_queries(text):
    
    llm = ChatGoogleGenerativeAI(
        model="gemini-3-flash-preview",
        temperature=0.3
    )

    query_prompt = PromptTemplate.from_template("""
You are a QA expert.

Generate ONLY 2 powerful queries:
1. Functional + validation scenarios
2. Edge cases + negative + boundary conditions

Make them detailed and non-overlapping.

Document:
{doc}
""")

    return (query_prompt | llm | StrOutputParser()).invoke({"doc": text})


# -------------------------------
# Core Function (RAG Pipeline)
# -------------------------------
def generate_test_cases(text, file_name):

    # -------------------------------
    #  Splitter
    # -------------------------------
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=400,
        chunk_overlap=80,
        separators=["\n\n", "\n", ".", " "]
    )

    chunks = splitter.split_text(text)

    docs = [
        LCDocument(page_content=chunk, metadata={"source": file_name})
        for chunk in chunks
    ]

    # -------------------------------
    # Embeddings
    # -------------------------------
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-2-preview"
    )

    # -------------------------------
    # LLM
    # -------------------------------
    llm = ChatGoogleGenerativeAI(
        model="gemini-3-flash-preview",
        temperature=0.3
    )

    # -------------------------------
    # Vector Store
    # -------------------------------
    vector_store = Chroma(
        persist_directory="chroma_db",
        embedding_function=embeddings
    )

    existing = vector_store.get(where={"source": file_name})

    if not existing or len(existing.get("documents", [])) == 0:
        ids = [str(uuid.uuid4()) for _ in docs]
        vector_store.add_documents(docs, ids=ids)

    # -------------------------------
    # Retriever (Simplified)
    # -------------------------------
    retriever = vector_store.as_retriever(
        search_type="similarity",
        search_kwargs={
            "k": 6,
            "filter": {"source": file_name}
        }
    )

    # -------------------------------
    # Query Generation (Cached)
    # -------------------------------
    queries_text = generate_queries(text)
    queries = [q.strip() for q in queries_text.split("\n") if q.strip()]

    combined_query = " OR ".join(queries)

    retrieved_docs = retriever.invoke(combined_query)

    unique_docs = list({doc.page_content: doc for doc in retrieved_docs}.values())

    def score(doc):
        return sum(q.lower() in doc.page_content.lower() for q in queries)

    ranked_docs = sorted(unique_docs, key=score, reverse=True)

    context = "\n\n".join(doc.page_content for doc in ranked_docs[:6])

    # -------------------------------
    # Final Prompt
    # -------------------------------
    prompt = PromptTemplate.from_template("""
You are a QA Test Case Generator.

STRICT REQUIREMENTS:
- Generate AT LEAST 15 test cases
- Cover:
  • Functional scenarios
  • Edge cases
  • Negative cases
  • Boundary conditions
  • Validation errors

FORMAT RULES:
1. Each test case must be clearly separated.
2. Use structured format.
3. Steps MUST be numbered and on separate lines.
4. Do NOT combine steps.
5. Each step must be short and actionable.

OUTPUT FORMAT:

Test Case ID: TC_001
Title: <Short Title>

Preconditions:
- <condition>

Steps:
1. Step one
2. Step two

Expected Result:
- Result

Priority: High/Medium/Low

-------------------------------------

REQUIREMENT DOCUMENT:
{context}
""")

    chain = prompt | llm | StrOutputParser()

    return chain.invoke({"context": context})


# -------------------------------
# File Reading Function
# -------------------------------
def extract_text(uploaded_file):
    text = ""

    if uploaded_file.name.endswith(".txt"):
        text = uploaded_file.read().decode("utf-8")

    elif uploaded_file.name.endswith(".docx"):
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text += " | ".join(row_text) + "\n"


    return text.strip()


# -------------------------------
# Streamlit UI
# -------------------------------
st.set_page_config(page_title="AI Test Case Generator", layout="wide")

st.title("🧪 AI Test Case Generator (RAG)")

uploaded_file = st.file_uploader("Upload Requirement Document", type=["txt", "docx"])

text = ""
file_name = ""

if uploaded_file is not None:
        st.success("✅ Document uploaded Successfully!")
        text = extract_text(uploaded_file)


# -------------------------------
# 🔹 Generate Button
# -------------------------------
if st.button("Generate Test Cases"):

    if not text.strip():
        st.warning("⚠️ Please upload a document")
    else:
        with st.spinner("Generating test cases..."):
            file_name = uploaded_file.name
            output = generate_test_cases(text , file_name)

            st.subheader("📋 Generated Test Cases")
            st.markdown(f"```\n{output}\n```")
